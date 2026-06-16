from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path("docs")
OUTPUT_PATH = OUTPUT_DIR / "Panduan_Penggunaan_E-Monev_KIP.docx"


ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(95, 99, 104)
LIGHT_FILL = "F5F7FA"
NOTE_FILL = "EEF5FF"
NOTE_BORDER = "BDD7EE"


@dataclass
class FigureItem:
    title: str
    description: str


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2F3", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Title", 24, ACCENT_DARK, 0, 12),
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "CaptionCustom" not in doc.styles:
        caption = doc.styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Calibri"
        caption.font.size = Pt(10)
        caption.font.italic = True
        caption.font.color.rgb = MUTED
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)

    if "NoteBox" not in doc.styles:
        note = doc.styles.add_style("NoteBox", WD_STYLE_TYPE.PARAGRAPH)
        note.font.name = "Calibri"
        note.font.size = Pt(10.5)
        note.font.color.rgb = TEXT
        note.paragraph_format.space_after = Pt(0)
        note.paragraph_format.line_spacing = 1.1


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Panduan Penggunaan\nE-Monev KIP")
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("Panduan operasional untuk pengguna dinas dan admin")
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_layout_fixed(meta)
    widths = [Cm(5), Cm(9.5)]
    data = [
        ("Nama Aplikasi", "E-Monev KIP"),
        ("Instansi", "Komisi Informasi Publik Kabupaten Banjarnegara"),
        ("Jenis Dokumen", "Panduan penggunaan aplikasi web"),
        ("Tanggal Penyusunan", "11 Juni 2026"),
    ]
    for row, (label, value) in zip(meta.rows, data):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, "D9E2F3")
            if idx == 0:
                set_cell_shading(cell, "EAF1FB")
            p_label = cell.paragraphs[0]
            p_label.paragraph_format.space_after = Pt(0)
            run = p_label.add_run(label if idx == 0 else value)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.bold = idx == 0
            run.font.color.rgb = TEXT

    doc.add_paragraph("")
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(12)
    summary.paragraph_format.space_after = Pt(0)
    rs = summary.add_run(
        "Dokumen ini disusun untuk membantu proses login, pengelolaan biodata, pengisian kuesioner, verifikasi, pelaporan, dan administrasi aplikasi."
    )
    rs.font.name = "Calibri"
    rs.font.size = Pt(11)
    rs.font.color.rgb = TEXT

    doc.add_page_break()


def add_note_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, NOTE_FILL)
    set_cell_border(cell, NOTE_BORDER, "10")
    p = cell.paragraphs[0]
    p.style = doc.styles["NoteBox"]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.bold = False


def add_figure_placeholder(doc: Document, number: int, item: FigureItem) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)
    cell = table.cell(0, 0)
    cell.width = Cm(15.8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell, "D0D7E2")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("[Tempatkan screenshot di sini]")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = MUTED

    caption = doc.add_paragraph(style="CaptionCustom")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(f"Gambar {number}. {item.title}")
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_before = Pt(0)
    desc.paragraph_format.space_after = Pt(8)
    rd = desc.add_run(item.description)
    rd.font.name = "Calibri"
    rd.font.size = Pt(10)
    rd.font.color.rgb = MUTED


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)


def add_section_intro(doc: Document, title: str, body: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    configure_styles(doc)
    add_cover(doc)

    doc.add_heading("Daftar Isi Singkat", level=1)
    add_bullets(
        doc,
        [
            "Bagian A: Gambaran aplikasi dan peran pengguna.",
            "Bagian B: Panduan pengguna dinas, mulai dari login sampai unduh laporan.",
            "Bagian C: Panduan admin, mulai dari pengelolaan data sampai pelaporan.",
            "Bagian D: Tips operasional dan catatan penggunaan screenshot.",
        ],
    )

    add_section_intro(
        doc,
        "1. Gambaran Umum Aplikasi",
        "E-Monev KIP adalah aplikasi monitoring dan evaluasi keterbukaan informasi publik. Aplikasi ini dipakai oleh dua kelompok utama, yaitu pengguna dinas/badan publik sebagai pengisi kuesioner dan admin sebagai pengelola, verifikator, serta penyusun laporan hasil penilaian.",
    )
    add_bullets(
        doc,
        [
            "Pengguna dinas memakai aplikasi untuk masuk, melengkapi biodata, mengisi kuesioner per kategori, memantau hasil penilaian, membaca notifikasi, dan mengunduh laporan PDF ketika hasil sudah terverifikasi.",
            "Admin memakai aplikasi untuk mengelola badan publik, kategori dan pertanyaan kuesioner, jadwal pelaksanaan, verifikasi jawaban, statistik, pengiriman pesan, dan laporan.",
            "Skor akhir dibentuk dari jawaban yang dinyatakan valid pada saat verifikasi admin.",
        ],
    )
    add_note_box(
        doc,
        "Saran penggunaan screenshot: pasang gambar tepat di bawah langkah yang relevan. Pertahankan urutan nomor gambar pada dokumen ini agar pembaca mudah mengikuti alur."
    )

    doc.add_heading("2. Peran dan Hak Akses", level=1)
    add_bullets(
        doc,
        [
            "Pengguna Dinas/Badan Publik: fokus pada pengisian data dan pemantauan hasil.",
            "Admin: fokus pada manajemen master data, validasi jawaban, analisis, dan pelaporan.",
        ],
    )

    add_section_intro(
        doc,
        "3. Panduan Pengguna Dinas",
        "Bagian ini menjelaskan alur kerja pengguna dinas dari awal masuk ke sistem sampai memperoleh hasil penilaian dan laporan akhir."
    )

    doc.add_heading("3.1 Login ke Aplikasi", level=2)
    add_steps(
        doc,
        [
            "Buka halaman login aplikasi E-Monev KIP.",
            "Masukkan alamat email yang terdaftar.",
            "Masukkan password akun pengguna.",
            "Aktifkan opsi Ingatkan Saya bila perangkat yang digunakan bersifat pribadi.",
            "Klik tombol Masuk untuk masuk ke dashboard pengguna.",
        ],
    )
    add_note_box(doc, "Jika pengguna lupa password, gunakan menu Lupa Password pada halaman login untuk memulai proses pemulihan akun.")
    add_figure_placeholder(
        doc,
        1,
        FigureItem("Halaman login pengguna", "Tampilkan form login yang berisi email, password, opsi Ingatkan Saya, dan tautan Lupa Password."),
    )

    doc.add_heading("3.2 Memahami Dashboard Pengguna", level=2)
    add_bullets(
        doc,
        [
            "Dashboard menampilkan ringkasan hasil penilaian untuk periode aktif atau periode terakhir.",
            "Informasi yang ditampilkan meliputi nilai akhir, klasifikasi/status, status verifikasi, dan nilai per kategori.",
            "Bagian bawah dashboard menampilkan data badan publik dan tombol Edit Biodata.",
            "Jika hasil sudah terverifikasi, tombol Unduh Laporan PDF akan tersedia.",
        ],
    )
    add_figure_placeholder(
        doc,
        2,
        FigureItem("Dashboard pengguna", "Tampilkan ringkasan hasil penilaian, status verifikasi, tabel nilai per kategori, dan kartu data badan publik."),
    )

    doc.add_heading("3.3 Mengubah Biodata Peserta", level=2)
    add_steps(
        doc,
        [
            "Pada dashboard, klik tombol Edit Biodata.",
            "Lengkapi atau perbarui data badan publik: nama, website, nomor telepon, email, dan alamat.",
            "Lengkapi atau perbarui data responden: nama responden, nomor telepon, dan jabatan.",
            "Klik Simpan Perubahan untuk menyimpan data terbaru.",
        ],
    )
    add_note_box(doc, "Pastikan data email, nomor telepon, dan alamat diperbarui dengan benar karena data ini akan dipakai sebagai referensi administrasi dan pelaporan.")
    add_figure_placeholder(
        doc,
        3,
        FigureItem("Form ubah biodata peserta", "Tampilkan halaman edit biodata yang memuat bagian Data Badan Publik dan Data Responden."),
    )

    doc.add_heading("3.4 Mengakses Informasi Kuesioner", level=2)
    add_steps(
        doc,
        [
            "Buka menu Kuesioner dari navigasi aplikasi.",
            "Periksa informasi nama responden, badan publik, nomor telepon, dan durasi/jadwal kuesioner.",
            "Jika jadwal aktif tersedia, klik tombol Mulai Mengisi.",
            "Jika tombol tidak aktif, periksa pesan status untuk mengetahui apakah jadwal belum dibuka atau pengisian sudah pernah diselesaikan.",
        ],
    )
    add_figure_placeholder(
        doc,
        4,
        FigureItem("Halaman informasi kuesioner", "Tampilkan identitas responden, badan publik, durasi kuesioner, serta tombol Mulai Mengisi."),
    )

    doc.add_heading("3.5 Mengisi Kuesioner", level=2)
    add_steps(
        doc,
        [
            "Pilih kategori kuesioner melalui tab kategori yang tersedia.",
            "Baca setiap pertanyaan dan definisi operasionalnya.",
            "Pilih jawaban Ya atau Tidak pada kolom Pilih Jawaban.",
            "Isi kolom Link Dokumen apabila pertanyaan meminta tautan dokumen pendukung.",
            "Unggah dokumen PDF apabila pertanyaan mensyaratkan file pendukung.",
            "Klik Simpan Jawaban setelah seluruh pertanyaan pada kategori aktif diisi.",
            "Lanjutkan proses yang sama hingga seluruh kategori selesai diisi.",
        ],
    )
    add_note_box(
        doc,
        "Dokumen pendukung yang diunggah harus berformat PDF dengan ukuran maksimal 20 MB per file. Pastikan link dokumen dapat diakses oleh pihak verifikator."
    )
    add_figure_placeholder(
        doc,
        5,
        FigureItem("Form pengisian kuesioner per kategori", "Tampilkan tabel pertanyaan berisi nomor, teks pertanyaan, definisi operasional, pilihan jawaban, link dokumen, dan upload dokumen."),
    )

    doc.add_heading("3.6 Membaca Notifikasi", level=2)
    add_steps(
        doc,
        [
            "Buka menu Notifikasi.",
            "Gunakan filter Semua, Belum Dibaca, atau Sudah Dibaca sesuai kebutuhan.",
            "Klik salah satu judul pesan untuk melihat isi lengkap notifikasi.",
            "Pesan yang dibuka akan otomatis ditandai sebagai sudah dibaca.",
        ],
    )
    add_figure_placeholder(
        doc,
        6,
        FigureItem("Halaman notifikasi pengguna", "Tampilkan daftar pesan di sisi kiri dan isi pesan aktif di sisi kanan."),
    )

    doc.add_heading("3.7 Melihat Hasil Penilaian dan Mengunduh Laporan", level=2)
    add_bullets(
        doc,
        [
            "Nilai akhir dan klasifikasi akan tampil pada dashboard.",
            "Status verifikasi menunjukkan apakah hasil masih menunggu tinjauan atau sudah terverifikasi.",
            "Jika admin memberi catatan validasi, catatan tersebut akan muncul pada bagian Catatan Validasi.",
            "Tombol Unduh Laporan PDF hanya muncul ketika hasil penilaian sudah berstatus Terverifikasi.",
        ],
    )
    add_figure_placeholder(
        doc,
        7,
        FigureItem("Dashboard dengan laporan tersedia", "Tampilkan kondisi ketika status verifikasi sudah selesai dan tombol Unduh Laporan PDF aktif."),
    )

    doc.add_heading("3.8 Keluar dari Aplikasi", level=2)
    add_steps(
        doc,
        [
            "Klik menu keluar/logout pada navigasi.",
            "Konfirmasikan tindakan keluar jika sistem menampilkan halaman konfirmasi.",
            "Pastikan sesi telah berakhir sebelum menutup browser, terutama jika perangkat digunakan bersama.",
        ],
    )

    add_section_intro(
        doc,
        "4. Panduan Admin",
        "Bagian ini menjelaskan fitur administrasi untuk pengelolaan data, verifikasi nilai, analisis statistik, pesan, dan laporan."
    )

    doc.add_heading("4.1 Login Admin", level=2)
    add_steps(
        doc,
        [
            "Buka halaman login admin pada alamat yang disediakan sistem.",
            "Masukkan email dan password admin.",
            "Klik tombol login untuk masuk ke panel admin.",
        ],
    )
    add_figure_placeholder(
        doc,
        8,
        FigureItem("Halaman login admin", "Tampilkan form login khusus admin sebelum masuk ke panel pengelolaan."),
    )

    doc.add_heading("4.2 Dashboard Admin", level=2)
    add_bullets(
        doc,
        [
            "Dashboard menampilkan jumlah PPID pelaksana terdaftar, dinas yang menunggu verifikasi, dinas terverifikasi, dan rekap hasil penilaian.",
            "Di bagian bawah terdapat daftar singkat verifikasi nilai dinas yang dapat langsung dibuka.",
            "Admin dapat memakai dashboard sebagai titik awal monitoring harian.",
        ],
    )
    add_figure_placeholder(
        doc,
        9,
        FigureItem("Dashboard admin", "Tampilkan kartu ringkasan statistik dan tabel List Verifikasi Nilai Dinas."),
    )

    doc.add_heading("4.3 Mengelola Data Badan Publik", level=2)
    add_steps(
        doc,
        [
            "Buka menu Badan Publik.",
            "Gunakan kotak pencarian untuk mencari nama badan publik, email, atau responden.",
            "Klik Tambah untuk membuat akun dinas baru beserta data badan publiknya.",
            "Klik Edit untuk memperbarui data akun atau profil badan publik.",
            "Klik Lihat Detail untuk melihat informasi badan publik lebih lengkap.",
            "Klik Hapus jika akun perlu dihapus dari sistem.",
        ],
    )
    add_note_box(doc, "Saat menambahkan akun baru, admin perlu menyiapkan nama responden, email login, password awal, data badan publik, nomor telepon, dan jabatan.")
    add_figure_placeholder(
        doc,
        10,
        FigureItem("Menu badan publik", "Tampilkan daftar badan publik beserta tombol Tambah, Edit, Lihat Detail, dan Hapus."),
    )

    doc.add_heading("4.4 Mengelola Kuesioner dan Kategori", level=2)
    add_steps(
        doc,
        [
            "Buka menu Kuesioner.",
            "Gunakan tombol Tambah Kategori untuk membuat kategori baru.",
            "Gunakan tombol Edit Kategori untuk memperbarui judul, nama, atau deskripsi kategori.",
            "Gunakan tombol Lihat Detail Pertanyaan untuk membuka daftar pertanyaan pada kategori terpilih.",
            "Gunakan tombol Hapus Kategori bila kategori sudah tidak dipakai.",
        ],
    )
    add_figure_placeholder(
        doc,
        11,
        FigureItem("Halaman kuesioner admin", "Tampilkan tab kategori, ringkasan kategori aktif, serta tombol Jadwal Kuesioner, Tambah Kategori, Edit, dan Detail Pertanyaan."),
    )

    doc.add_heading("4.5 Mengatur Jadwal Kuesioner", level=2)
    add_steps(
        doc,
        [
            "Dari menu Kuesioner, klik Jadwal Kuesioner.",
            "Tambahkan atau perbarui periode pelaksanaan kuesioner sesuai tahun dan rentang tanggal.",
            "Pastikan jadwal yang diaktifkan sesuai periode pelaksanaan agar pengguna dapat mengisi kuesioner pada waktu yang tepat.",
        ],
    )
    add_figure_placeholder(
        doc,
        12,
        FigureItem("Menu jadwal kuesioner", "Tampilkan halaman pengelolaan jadwal yang memuat periode aktif atau daftar jadwal yang tersedia."),
    )

    doc.add_heading("4.6 Mengelola Pertanyaan Kuesioner", level=2)
    add_bullets(
        doc,
        [
            "Pertanyaan disusun per kategori dan dikaitkan dengan jadwal tertentu.",
            "Admin dapat menentukan urutan pertanyaan, definisi operasional, skor maksimum, serta kebutuhan link atau upload dokumen.",
            "Gunakan halaman detail pertanyaan untuk meninjau daftar pertanyaan yang berlaku pada kategori dan jadwal tertentu.",
        ],
    )
    add_figure_placeholder(
        doc,
        13,
        FigureItem("Detail pertanyaan kategori", "Tampilkan daftar pertanyaan pada salah satu kategori lengkap dengan aksi tambah atau edit pertanyaan."),
    )

    doc.add_heading("4.7 Verifikasi Nilai Dinas", level=2)
    add_steps(
        doc,
        [
            "Buka daftar verifikasi dari dashboard admin atau menu Penilaian.",
            "Pilih salah satu dinas dan periode jadwal yang akan diverifikasi.",
            "Tinjau jawaban, link dokumen, dan file pendukung pada setiap pertanyaan.",
            "Tentukan status Valid atau Tidak Valid untuk setiap jawaban.",
            "Tambahkan catatan validasi bila diperlukan sebagai umpan balik.",
            "Klik Simpan Progress apabila verifikasi belum selesai.",
            "Klik Selesai Verifikasi jika seluruh pertanyaan sudah ditinjau.",
        ],
    )
    add_note_box(
        doc,
        "Nilai kategori dan nilai akhir dihitung dari jawaban yang berstatus Valid dan bernilai Ya. Semua pertanyaan harus divalidasi sebelum proses verifikasi dapat diselesaikan."
    )
    add_figure_placeholder(
        doc,
        14,
        FigureItem("Halaman verifikasi nilai dinas", "Tampilkan kartu ringkasan nilai, progres validasi, daftar pertanyaan, tombol Valid/Tidak Valid, area catatan, dan tombol Selesai Verifikasi."),
    )

    doc.add_heading("4.8 Melihat Statistik", level=2)
    add_bullets(
        doc,
        [
            "Menu Statistik menampilkan progres verifikasi, rata-rata nilai per kategori, distribusi klasifikasi, 10 badan publik teratas, 10 badan publik terbawah, statistik per pertanyaan, dan tren tahunan.",
            "Admin dapat mengganti periode jadwal untuk melihat statistik pada tahun atau periode tertentu.",
            "Grafik membantu analisis cepat, sedangkan tabel di bawah grafik memudahkan pengecekan rinci.",
        ],
    )
    add_figure_placeholder(
        doc,
        15,
        FigureItem("Halaman statistik", "Tampilkan pemilih periode, kartu progres verifikasi, dan salah satu tampilan grafik statistik."),
    )

    doc.add_heading("4.9 Mengirim Pesan Notifikasi", level=2)
    add_steps(
        doc,
        [
            "Buka menu Pesan.",
            "Tinjau daftar pesan yang sudah pernah dibuat apabila diperlukan.",
            "Isi judul notifikasi dan isi pesan.",
            "Pilih channel pengiriman: aplikasi, email, atau keduanya.",
            "Pilih target penerima: semua akun dinas, akun tertentu, atau akun berdasarkan kondisi.",
            "Klik Kirim untuk mengirim notifikasi.",
        ],
    )
    add_bullets(
        doc,
        [
            "Kondisi target dapat dipilih untuk akun yang belum mengisi, sudah mengisi, menunggu verifikasi, atau sudah terverifikasi.",
            "Jika channel email dipilih, sistem akan mencoba mengirim email saat itu juga.",
        ],
    )
    add_figure_placeholder(
        doc,
        16,
        FigureItem("Menu pesan admin", "Tampilkan daftar pesan di bagian atas dan form pengiriman notifikasi di bagian bawah."),
    )

    doc.add_heading("4.10 Mengunduh Laporan", level=2)
    add_steps(
        doc,
        [
            "Buka menu Laporan.",
            "Gunakan pencarian atau filter jadwal untuk menemukan badan publik yang diinginkan.",
            "Periksa nilai akhir dan klasifikasi yang sudah terverifikasi.",
            "Klik tombol Unduh untuk menghasilkan laporan PDF per badan publik.",
        ],
    )
    add_note_box(doc, "Laporan hanya tersedia untuk hasil yang sudah berstatus Terverifikasi.")
    add_figure_placeholder(
        doc,
        17,
        FigureItem("Halaman laporan admin", "Tampilkan daftar badan publik, filter jadwal, nilai akhir, klasifikasi, dan tombol Unduh."),
    )

    doc.add_heading("4.11 Mengubah Pengaturan Profil Admin", level=2)
    add_steps(
        doc,
        [
            "Buka menu Pengaturan.",
            "Perbarui nama, email, password, atau foto profil admin.",
            "Klik Simpan untuk menyimpan perubahan.",
        ],
    )
    add_figure_placeholder(
        doc,
        18,
        FigureItem("Halaman pengaturan admin", "Tampilkan form profil admin yang memuat nama, email, password baru, dan foto profil."),
    )

    add_section_intro(
        doc,
        "5. Tips Operasional",
        "Bagian ini dapat digunakan sebagai ringkasan praktik yang disarankan agar penggunaan aplikasi berjalan lebih lancar."
    )
    add_bullets(
        doc,
        [
            "Lakukan pembaruan biodata sebelum jadwal pengisian dibuka untuk menghindari kendala administrasi.",
            "Siapkan dokumen pendukung dalam format PDF sebelum mulai mengisi kuesioner.",
            "Gunakan nama file dokumen yang jelas agar memudahkan proses verifikasi.",
            "Admin disarankan menyimpan progress verifikasi secara berkala ketika memeriksa banyak pertanyaan.",
            "Gunakan fitur pesan untuk mengingatkan akun dinas yang belum mengisi atau masih menunggu verifikasi.",
        ],
    )

    doc.add_heading("6. Daftar Keterangan Gambar", level=1)
    figure_titles = [
        "Halaman login pengguna",
        "Dashboard pengguna",
        "Form ubah biodata peserta",
        "Halaman informasi kuesioner",
        "Form pengisian kuesioner per kategori",
        "Halaman notifikasi pengguna",
        "Dashboard dengan laporan tersedia",
        "Halaman login admin",
        "Dashboard admin",
        "Menu badan publik",
        "Halaman kuesioner admin",
        "Menu jadwal kuesioner",
        "Detail pertanyaan kategori",
        "Halaman verifikasi nilai dinas",
        "Halaman statistik",
        "Menu pesan admin",
        "Halaman laporan admin",
        "Halaman pengaturan admin",
    ]
    for idx, title in enumerate(figure_titles, start=1):
        doc.add_paragraph(f"Gambar {idx}. {title}")

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
